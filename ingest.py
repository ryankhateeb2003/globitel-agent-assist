import sys
import io
from pathlib import Path
import json

from app.ingestion.pipeline import extract_document
from app.ingestion.html_extractor import extract_html
from app.ingestion.docx_extractor import extract_docx
from app.ingestion.pdf_extractor import extract_pdf

# Ensure UTF-8 output encoding for terminal
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")


def get_raw_extraction(file_path: Path) -> str:
    """
    Get raw text from document before generic cleaning and normalisation.
    """
    suffix = file_path.suffix.lower()

    if suffix in {".html", ".htm"}:
        html = file_path.read_text(encoding="utf-8", errors="replace")
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text("\n", strip=True)

    elif suffix == ".docx":
        from docx import Document

        doc = Document(file_path)
        lines = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                lines.append(
                    " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                )
        return "\n".join(lines)

    elif suffix == ".pdf":
        import pymupdf

        lines = []
        with pymupdf.open(file_path) as doc:
            for page in doc:
                lines.append(page.get_text())
        return "\n".join(lines)

    return ""


def run_ingestion(corpus_dir: str | Path = "corpus"):
    corpus_path = Path(corpus_dir)

    if not corpus_path.exists():
        raise FileNotFoundError(f"Corpus directory not found: {corpus_path}")

    samples_dir = Path("extraction-samples")
    samples_dir.mkdir(exist_ok=True)

    processed_docs = []
    ignored_files = []

    # Sample files to record raw vs cleaned output for Task 1 deliverable
    sample_targets = {
        "corpus/en/help-center.html": "help-center_en_html.txt",
        "corpus/en/orange-money.docx": "orange-money_en_docx.txt",
        "corpus/en/mobile-lines.pdf": "mobile-lines_en_pdf.txt",
        "corpus/ar/help-center.html": "help-center_ar_html.txt",
        "corpus/ar/orange-money.docx": "orange-money_ar_docx.txt",
        "corpus/ar/mobile-lines.pdf": "mobile-lines_ar_pdf.txt",
    }

    files = sorted(corpus_path.glob("**/*.*"))

    for file_path in files:
        if file_path.name.startswith(".") or "empty" in file_path.name or "scanned" in file_path.name:
            ignored_files.append(str(file_path))
            continue

        try:
            cleaned_text = extract_document(file_path)

            doc_entry = {
                "file_path": str(file_path.as_posix()),
                "language": "ar" if "/ar/" in str(file_path.as_posix()) else "en",
                "format": file_path.suffix.lower().lstrip("."),
                "char_count": len(cleaned_text),
                "line_count": len(cleaned_text.splitlines()),
            }
            processed_docs.append(doc_entry)

            # Check if this file is one of the 6 extraction samples
            posix_path = str(file_path.as_posix())
            for target_path, sample_filename in sample_targets.items():
                if target_path in posix_path:
                    raw_text = get_raw_extraction(file_path)
                    sample_file = samples_dir / sample_filename
                    sample_content = (
                        f"=== RAW EXTRACTION ({file_path.name}) ===\n"
                        f"{raw_text[:2000]}\n\n"
                        f"===================================================\n"
                        f"=== CLEANED & NORMALISED EXTRACTION ===\n"
                        f"===================================================\n"
                        f"{cleaned_text[:2000]}\n"
                    )
                    sample_file.write_text(sample_content, encoding="utf-8")
                    print(f"[SAMPLE CREATED] {sample_filename}")

        except Exception as exc:
            print(f"[REJECTED] {file_path}: {exc}")

    print(f"\n================ INGESTION SUMMARY ================")
    print(f"Successfully processed documents: {len(processed_docs)}")
    print(f"Ignored / Rejected documents    : {len(ignored_files)}")

    return processed_docs


if __name__ == "__main__":
    run_ingestion()
