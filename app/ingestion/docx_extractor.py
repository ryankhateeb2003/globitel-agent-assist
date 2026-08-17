from pathlib import Path

from docx import Document


# =========================================================
# Exact website/UI elements that should not enter the RAG corpus
# =========================================================

REMOVE_EXACT = {
    "Skip to main content",
    "Free Delivery",
    "Main navigation",
    "Personal",
    "Business",
    "Corporate",
    "English",
    "العربية",
    "العربيه",
    "Accessibility",
    "switch",
    "Dark mode",
    "A-",
    "A",
    "A+",
    "Font size",
    "Account",
    "Home",
    "Apps Menu",
    "Contact Us",
    "Contact us",
    "Messenger",
    "Instagram",
    "Whatsapp",
    "Shops & Coverage",
    "Find Shops",
    "Internet Coverage",
    "Quick Pay / Refill",
    "Search",
    "Useful FAQs",
    "Useful Links",
    "Helpful short codes",
    "Bill Payment",
    "Internet Devices User guides",
    "Orange products",
    "- All -",
    "Apply",
    "Was this page helpful ?",
    "Submit",
    "Thank you for your feedback",
    "Report a problem",
    "Still need more help ?",
    "Get in touch with us and we're happy to respond to all your queries",
    "How can we help you ?",
    "Check out useful information",
    "FAQ & Help",
    "Visit our Store",
    "Footer",
    "Fixed Lines",
    "Small Business",
    "Enterprise",
    "Business eShop",
    "About Orange",
    "Orange CSR",
    "Investors Relations",
    "Media Center",
    "Careers",
    "Wholesale",
    "Compliance & Fraud",
    "Distributor's corner",
    "Legal",
    "Orange Max it",
    "We Accept",
    "Join our newsletter",
    "to receive Offers & Promotions",
    "Email",
    "Subscribe",
    "Privacy policy",
    "Site map",
    "Blog",
    "Terms & conditions",
    "Return Policy",
}


# =========================================================
# Navigation text that should not enter the RAG corpus
# =========================================================

NAVIGATION_BLOCK = {
    "Orange Money Help",
    "English العربية",
    "Internet",
    "Internet Offers",
    "All",
    "Fiber Offers",
    "5G Home Internet",
    "4G Flybox Home Internet",
    "Prepaid Orange Internet",
    "Visitors Lines",
    "Satellite Offers",
    "ADSL Offers",
    "Daman Offers",
    "Internet Services",
    "OSN",
    "TOD",
    "Anghami",
    "Gaming Control",
    "More Services",
    "Max it App",
    "Orange Money",
    "Mobile Lines",
    "Mobile Offers",
    "All mobile offers",
    "Ma'ak Lines",
    "Humat Al- Watan max",
    "Mobile Services",
    "International",
    "Roaming",
    "e-Sh7anli",
    "Devices & Accessories",
    "All Devices",
    "Accessories",
    "Wearables",
    "(Watches, Headsets & More)",
    "Home Entertainment",
    "(TVs, Speakers, Gaming...)",
    "Tablets & Laptops",
    "SmartLife",
    "Network Devices",
    "Mobile Brands",
    "Apple",
    "Samsung",
    "Xiaomi",
    "Discover Devices",
    "Max it",
    "All Max it Offers",
    "My Line",
    "MarketPlace",
    "Rewards",
    "ِAbout Orange Money",
    "Our Services",
    "Self Registration",
    "Our Discounts",
    "Transaction Fees and Limits",
    "FAQs",
    "Help",
    "Frequently Asked Questions",
    "All FAQs",
    "Fiber and ADSL",
    "eShop",
    "International and Roaming",
    "Payment Methods",
    "More Topics",
    "Help Center How can we help ?",
    "Help Center",
    "Orange Money - Help",
    "Find help answers on Orange Money Service, visit our help center to get more information on using the Mobile Wallet.",
    "Getting started",
    "Facing a problem",
    "- Any -",
    "Facing a problem (eShop)",
    "Getting started (eShop)",
    "Prepaid",
    "Postpaid",
    "ADSL",
    "Fiber",
    "Getting started (OM)",
    "Facing a problem (OM)",
    "Bills Payment",
    "Fixed Line",
    "General FAQs",
    "About ADSL",
    "Call Waiting",
    "QR FAQs",
    "CliQ FAQs",
    "International transfers",
    "Local Transfer",
    "e-Voucher FAQs",
    "Google Pay",
    "Card to wallet Top up",
    "Virtual Visa Card",
    "Visa Cards",
    "- All -General FAQsAbout ADSLCall WaitingQR FAQsCliQ FAQsInternational transfersLocal Transfere-Voucher FAQsGoogle PayCard to wallet Top upVirtual Visa CardVisa Cards",
}


def remove_duplicate_lines(lines: list[str]) -> list[str]:
    """
    Remove exact duplicate lines while preserving order.
    """

    seen = set()
    result = []

    for line in lines:
        line = line.strip()

        if not line:
            continue

        if line in seen:
            continue

        seen.add(line)
        result.append(line)

    return result


def should_remove(text: str) -> bool:
    """
    Return True when a paragraph is known general UI, navigation,
    or footer boilerplate.
    """

    if text in REMOVE_EXACT:
        return True

    if text in NAVIGATION_BLOCK:
        return True

    lower_text = text.lower()

    if "have a new question? post your question" in lower_text:
        return True

    if "thank you for your feedback" in lower_text:
        return True

    if "report a problem" in lower_text:
        return True

    if "still need more help" in lower_text:
        return True

    if "having an issue with one of our services" in lower_text:
        return True

    if "join our newsletter" in lower_text:
        return True

    if "all rights reserved" in lower_text or "جميع الحقوق محفوظة" in lower_text:
        return True

    return False


def extract_docx(file_path: str | Path) -> str:
    """
    Extract clean documentation text from a DOCX file using general rules.

    Processing:
    1. Validate the file.
    2. Read DOCX paragraphs.
    3. Filter out generic website UI, navigation menus, and footer boilerplate.
    4. Remove exact duplicate lines.
    5. Return clean documentation text.
    """

    file_path = Path(file_path)

    # -----------------------------------------------------
    # 1. Validate file
    # -----------------------------------------------------

    if not file_path.exists():
        raise FileNotFoundError(
            f"DOCX file not found: {file_path}"
        )

    if file_path.stat().st_size == 0:
        raise ValueError(
            f"DOCX file is empty: {file_path}"
        )

    # -----------------------------------------------------
    # 2. Read DOCX
    # -----------------------------------------------------

    try:
        document = Document(file_path)
    except Exception as exc:
        raise ValueError(
            f"Could not read DOCX file: {file_path}"
        ) from exc

    raw_lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            raw_lines.append(text)

    for table in document.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                raw_lines.append(" | ".join(row_cells))

    if not raw_lines:
        raise ValueError(
            f"DOCX contains no usable content: {file_path}"
        )

    # -----------------------------------------------------
    # 3. Filter generic boilerplate & UI noise
    # -----------------------------------------------------

    cleaned = []

    for line in raw_lines:
        line = line.strip()

        if not line:
            continue

        if line == "Footer" or line == "التذييل":
            continue

        if should_remove(line):
            continue

        cleaned.append(line)

    # -----------------------------------------------------
    # 4. Remove exact duplicate lines
    # -----------------------------------------------------

    cleaned = remove_duplicate_lines(cleaned)

    # -----------------------------------------------------
    # 5. Build final text
    # -----------------------------------------------------

    clean_text = "\n".join(cleaned)

    if not clean_text:
        raise ValueError(
            f"DOCX extraction produced no usable content: {file_path}"
        )

    return clean_text