"""
Generates a Word document summarizing every question type/case this
system (Tasks 4-6) handles, with example questions and expected
behaviour -- a readable reference table, not a new deliverable required
by the task brief. Requires python-docx (already a project dependency).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROWS = [
    # (case, example_question, expected_behavior, status)
    ("Normal question, answer in docs",
     "What is QR Payment?",
     "Direct grounded answer + sources list. Answered in the question's language.",
     "Tested"),
    ("Same, in Arabic",
     "شو رسوم استخدام Orange Money؟",
     "Answers in Arabic, even if the source chunk is English (cross-language retrieval, Task 3 decision).",
     "Verified in Task 3/4"),
    ("Exact value (short code / fee)",
     "شو الكود يلي بحدد فيه حد التجوال؟",
     "Hybrid mode (vector+keyword+rerank) catches exact codes/figures vector search alone tends to miss (Task 5).",
     "Verified in Task 5"),
    ("In-domain but not documented",
     "Do you offer discounted family postpaid bundles?",
     'Refuses clearly: "I don\'t have this information in the documentation..." No guessing.',
     "Tested"),
    ("Outside the domain entirely",
     "What's the weather like in Amman today?",
     'Refuses immediately: "I can only help with questions about Orange Jordan\'s..." Refuses before generating an answer.',
     "Tested"),
    ("Ambiguous question",
     "كيف بلغيه؟",
     "Asks ONE clarifying question naming the plausible services, instead of guessing.",
     "Tested"),
    ("Needs customer account data",
     "Why was I charged 5 JOD last month?",
     "Refuses and tells the agent to check the account system, does not fabricate a number.",
     "Tested"),
    ("Sources conflict",
     "(constructed) What's the max wallet-to-wallet transfer amount?",
     "Presents both conflicting figures and names each source. Does not silently pick one.",
     "Tested"),
    ("Levantine dialect (Arabic script)",
     "كيف بقدر اعبي محفظتي؟",
     "Guardrail must still answer (not refuse for being dialect). Known gap: informal wording can still miss the right chunk at retrieval.",
     "Tested -- partial gap found"),
    ("Arabizi (Latin script)",
     "kif ba3mal top up la mahfazti",
     "Detected, transliterated to Arabic for retrieval, and answered normally -- not refused.",
     "Tested -- bug found & fixed"),
    ("Empty question",
     '""',
     "400 Bad Request, rejected before any retrieval or Groq call.",
     "Verified in Task 4"),
    ("Very long question (>1000 chars)",
     "(1199-character repeated string)",
     "400 Bad Request with the actual length and the limit stated.",
     "Verified in Task 4"),
    ("Mixed-language question",
     "شو رسوم استخدام QR Payment؟",
     "Detects the dominant language correctly despite an embedded English term, answers in that language.",
     "Verified in Task 4"),
]

STATUS_COLORS = {
    "Tested": "1E7C34",
    "Tested -- partial gap found": "B26A00",
    "Tested -- bug found & fixed": "1E7C34",
    "Verified in Task 3/4": "555555",
    "Verified in Task 4": "555555",
    "Verified in Task 5": "555555",
}

HEADER_BG = "F58220"  # Orange brand-ish accent for the header row


def set_cell_background(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def build():
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    title = doc.add_heading("Task 6 -- Question Coverage Reference", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    intro = doc.add_paragraph()
    intro.add_run(
        "Every question type the Globitel bilingual agent-assist system "
        "handles, with a representative example and the expected "
        "behaviour, drawn from Tasks 4-6. Rows marked \"Tested\" were "
        "verified with a real /ask request during this session; rows "
        "marked \"Verified in Task N\" were established and documented in "
        "that earlier task's deliverables."
    ).italic = True

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    widths = [Inches(1.7), Inches(2.3), Inches(3.0), Inches(1.3)]
    headers = ["Case", "Example Question", "Expected Behaviour", "Status"]

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].width = widths[i]
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        set_cell_background(hdr_cells[i], HEADER_BG)

    for case, question, behavior, status in ROWS:
        row_cells = table.add_row().cells
        for i, val in enumerate([case, question, behavior, status]):
            row_cells[i].width = widths[i]
            p = row_cells[i].paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            if i == 3:
                run.bold = True
                color_hex = STATUS_COLORS.get(status, "000000")
                run.font.color.rgb = RGBColor.from_string(color_hex)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Note: ").bold = True
    note.add_run(
        "the dialect (Arabic-script) row is the one open, documented "
        "limitation -- the guardrail layer correctly does not refuse for "
        "dialect phrasing, but retrieval can still miss the right chunk "
        "when the dialect word doesn't lexically match the corpus's "
        "formal wording (e.g. \"اعبي\" vs \"اشحن\"). See dialect-tests.md "
        "for the full writeup."
    )

    out_path = "app/guardrails/task6-question-coverage.docx"
    doc.save(out_path)
    print(f"[SAVED] {out_path}")


if __name__ == "__main__":
    build()
